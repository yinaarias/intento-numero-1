import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO

st.title("Generador de Autos de Reprogramación Judicial")

uploaded_file = st.file_uploader("Sube el archivo Excel con los datos", type=["xlsx"])

if uploaded_file:
    df = pd.read_excel(uploaded_file, engine="openpyxl")
    st.success(f"{len(df)} registros cargados correctamente.")

    for index, row in df.iterrows():
        doc = Document()
        doc.add_heading("AUTO DE REPROGRAMACIÓN DE AUDIENCIA", level=1)

        doc.add_paragraph(f"C.U.I.: {row['CUI']}")
        doc.add_paragraph(f"N.I.: {row['NI']}")
        doc.add_paragraph(f"Tipo de Procedimiento: {row['Tipo Procedimiento']}")
        doc.add_paragraph(f"Acusado: {row['Acusado']}")
        doc.add_paragraph(f"Delito: {row['Delito']}")
        doc.add_paragraph(f"Fecha del Auto: {row['Fecha Auto']}")
        doc.add_paragraph(f"Juez: {row['Juez']}")

        doc.add_paragraph("\nVista la constancia que antecede, este Despacho Judicial REPROGRAMA la diligencia y, "
                          "en consecuencia, se fija fecha y hora para AUDIENCIA CONCENTRADA dentro del proceso de la referencia, "
                          f"el día {row['Nueva Fecha Audiencia']}, a las {row['Hora Audiencia']} horas.\n\nCúmplase.")

        doc.add_paragraph(f"\n{row['Juez']}\nJuez")

        doc.add_paragraph("\n---\n")
        doc.add_paragraph("CONSTANCIA SECRETARIAL")
        doc.add_paragraph(row["Constancia Secretarial"])

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.download_button(
            label=f"Descargar Auto {index + 1}",
            data=buffer,
            file_name=f"Auto_Reprogramacion_{index + 1}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
